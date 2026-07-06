from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_PATH = Path(r"C:\Users\aless\Desktop\Progetto_FP\docs\StoreHub_Contratto_SaaS_Draft.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D9E2F3")


def set_run_font(run, name="Calibri", size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=6, line=1.10):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    style_paragraph(p, before=0, after=8, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=20, bold=True, color="1F3A5F")
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=0, after=10, line=1.0)
    run = p.add_run(text)
    set_run_font(run, size=10, color="5B6B8C")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(
        p,
        before=16 if level == 1 else 10,
        after=8 if level == 1 else 6,
        line=1.0,
    )
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=15, bold=True, color="2E74B5")
    elif level == 2:
        set_run_font(run, size=12.5, bold=True, color="2E74B5")
    else:
        set_run_font(run, size=11.5, bold=True, color="1F4D78")
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    style_paragraph(p)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style_paragraph(p, before=0, after=4, line=1.15)
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        style_paragraph(p, before=0, after=4, line=1.15)
        run = p.add_run(item)
        set_run_font(run)


def add_info_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.6)
    for left, right in rows:
        row = table.add_row()
        row.cells[0].width = Inches(1.9)
        row.cells[1].width = Inches(4.6)
        for idx, value in enumerate((left, right)):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            style_paragraph(p, before=0, after=0, line=1.05)
            run = p.add_run(value)
            set_run_font(run, bold=(idx == 0))
            if idx == 0:
                set_cell_shading(cell, "F2F4F7")
    set_table_borders(table)
    return table


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, before=0, after=0, line=1.0)
    run = p.add_run("StoreHub - bozza contrattuale")
    set_run_font(run, size=9, color="6B7280")


def add_page_break_section(doc):
    doc.add_section(WD_SECTION.NEW_PAGE)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_footer(section)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(11)

    add_title(doc, "Contratto di concessione in uso del software StoreHub")
    add_subtitle(
        doc,
        "Bozza operativa da condividere con consulente legale, privacy e amministrazione.",
    )

    add_info_table(
        doc,
        [
            ("Fornitore", "[Ragione sociale del concedente]"),
            ("Cliente", "[Ragione sociale del concessionario / licenziatario]"),
            ("Data", "[gg/mm/aaaa]"),
            ("Versione documento", "Draft 1 - predisposto su base architettura attuale di StoreHub"),
        ],
    )

    add_heading(doc, "1. Parti e premessa", 1)
    add_body(
        doc,
        "Il presente contratto disciplina la concessione in uso, in modalità software-as-a-service (SaaS), della piattaforma StoreHub e dei relativi moduli applicativi, come meglio descritti nei successivi articoli e negli allegati.",
    )
    add_body(
        doc,
        "Le premesse, gli allegati e gli eventuali ordini/appendici economiche costituiscono parte integrante e sostanziale del presente accordo.",
    )

    add_heading(doc, "2. Oggetto", 1)
    add_body(
        doc,
        "Il Fornitore concede al Cliente un diritto non esclusivo, non cedibile e non sublicenziabile di utilizzo della piattaforma StoreHub, accessibile via web, per la gestione operativa e di reporting dei punti vendita, dei relativi utenti e delle configurazioni tenant abilitate.",
    )
    add_bullets(
        doc,
        [
            "dashboard e cruscotti operativi;",
            "magazzino, ordini fornitore e listini;",
            "rendiconto e distinta cassa;",
            "gestione orari e anagrafiche correlate;",
            "link operativi, estrazioni ed eventuali moduli opzionali;",
            "funzioni di amministrazione tenant e, se previste, funzioni master/piattaforma.",
        ],
    )
    add_body(
        doc,
        "La concessione riguarda l'uso del software e non comporta in alcun modo il trasferimento della proprietà del codice sorgente, della documentazione interna di sviluppo o dei diritti di proprietà intellettuale.",
    )

    add_heading(doc, "3. Natura del servizio", 1)
    add_bullets(
        doc,
        [
            "StoreHub è erogato come applicazione web multi-tenant.",
            "L'accesso avviene tramite credenziali utente e profili di autorizzazione configurati dal Fornitore e/o dagli amministratori del tenant.",
            "Il servizio può comprendere moduli standard, moduli opzionali e personalizzazioni specifiche per singolo tenant.",
        ],
    )

    add_heading(doc, "4. Architettura generale e moduli", 1)
    add_body(
        doc,
        "Alla data della presente bozza, la piattaforma è basata su backend Python/Flask, pubblicazione web su infrastruttura Azure App Service, basi dati SQL Server per la componente applicativa e componenti Supabase per autenticazione, profili, alcune configurazioni e servizi complementari, come meglio dettagliato nell'Allegato Tecnico.",
    )
    add_body(
        doc,
        "Le parti prendono atto che alcuni tenant possono avere moduli disattivati, configurazioni dedicate, database applicativi separati o integrazioni esterne opzionali.",
    )

    add_heading(doc, "5. Utenti, ruoli e tenant", 1)
    add_bullets(
        doc,
        [
            "Il Cliente utilizza il software all'interno del/i tenant a lui assegnato/i.",
            "Gli utenti possono essere profilati, a titolo esemplificativo, come user, admin tenant, master o altri ruoli applicativi configurati.",
            "Il Fornitore può definire limiti per numero utenti, numero store, moduli attivi, funzionalità abilitate e capacità operative del tenant.",
            "Il Cliente è responsabile della correttezza delle abilitazioni richieste per i propri utenti e del rispetto delle regole di utilizzo.",
        ],
    )

    add_heading(doc, "6. Corrispettivi", 1)
    add_body(
        doc,
        "I corrispettivi per attivazione, canone ricorrente, eventuali personalizzazioni, supporto evolutivo, migrazioni, import storici, moduli aggiuntivi o servizi professionali saranno disciplinati nell'offerta economica o in appendici commerciali richiamate dal presente contratto.",
    )

    add_heading(doc, "7. Livelli di servizio e supporto", 1)
    add_bullets(
        doc,
        [
            "Il Fornitore assicura la messa a disposizione del servizio con diligenza professionale e secondo le caratteristiche tecniche pattuite.",
            "Eventuali SLA di disponibilità, tempi di presa in carico, tempi di ripristino, finestre di manutenzione e canali di supporto devono essere indicati in apposito allegato o appendice commerciale.",
            "Sono esclusi dai livelli di servizio i disservizi causati da fornitori terzi, reti del Cliente, browser obsoleti, errori di configurazione del Cliente o utilizzi non conformi.",
        ],
    )

    add_heading(doc, "8. Aggiornamenti, manutenzione ed evoluzione", 1)
    add_body(
        doc,
        "Il Fornitore potrà effettuare aggiornamenti correttivi, adeguativi, migliorativi o evolutivi della piattaforma. Gli aggiornamenti che incidono sul perimetro funzionale, sui flussi operativi o sui costi saranno previamente concordati se esulano dal servizio ordinario incluso.",
    )

    add_heading(doc, "9. Integrazioni di terze parti", 1)
    add_bullets(
        doc,
        [
            "La piattaforma può integrarsi con servizi terzi, inclusi ma non limitati a provider POS, sistemi Microsoft/SharePoint, servizi AI, piattaforme survey, sistemi di recensioni, servizi di autenticazione o API di fornitori esterni.",
            "Il Cliente prende atto che disponibilità, continuità, costi e termini d'uso dei servizi terzi dipendono dai rispettivi fornitori.",
            "Eventuali chiavi API, token, client secret o credenziali di integrazione devono essere custoditi e gestiti secondo procedure sicure concordate tra le parti.",
        ],
    )

    add_heading(doc, "10. Proprietà intellettuale", 1)
    add_body(
        doc,
        "Tutti i diritti di proprietà intellettuale e industriale relativi al software StoreHub, ai suoi sviluppi, personalizzazioni riutilizzabili, librerie, modelli, documentazione tecnica e organizzazione del codice restano di esclusiva titolarità del Fornitore, salvo diverso accordo scritto.",
    )

    add_heading(doc, "11. Obblighi del Cliente", 1)
    add_numbered(
        doc,
        [
            "utilizzare il servizio conformemente alla documentazione funzionale e alle istruzioni ricevute;",
            "custodire con cura le credenziali di accesso e limitarne l'uso ai soggetti autorizzati;",
            "non tentare di accedere al codice sorgente, aggirare misure di sicurezza, estrarre dati di altri tenant o utilizzare il servizio in modo illecito;",
            "fornire dati corretti, leciti e pertinenti ai fini del trattamento e dell'operatività del servizio;",
            "designare internamente i referenti autorizzati a richiedere configurazioni, attivazioni o modifiche sensibili.",
        ],
    )

    add_heading(doc, "12. Protezione dei dati personali", 1)
    add_body(
        doc,
        "Qualora il Fornitore tratti dati personali per conto del Cliente nell'ambito dell'erogazione del servizio, le parti disciplineranno tale rapporto nell'Allegato Privacy / Data Processing Agreement allegato al presente contratto.",
    )
    add_body(
        doc,
        "Il Cliente resta titolare del trattamento dei dati caricati e trattati tramite la piattaforma per le finalità proprie della propria organizzazione, salvo diversa qualificazione espressamente pattuita.",
    )

    add_heading(doc, "13. Riservatezza", 1)
    add_body(
        doc,
        "Le parti si impegnano a mantenere riservate le informazioni tecniche, commerciali, organizzative e di sicurezza apprese in occasione dell'esecuzione del contratto, incluse architetture, credenziali, configurazioni, chiavi API, logiche di integrazione e dati non pubblici.",
    )

    add_heading(doc, "14. Limitazioni di responsabilità", 1)
    add_body(
        doc,
        "Salvo dolo o colpa grave e salvo i limiti inderogabili di legge, il Fornitore non sarà responsabile per danni indiretti, perdita di profitto, perdita di chance, perdita di dati imputabile a condotte del Cliente o di terzi, sospensioni di servizi di terze parti, malfunzionamenti dovuti a input errati o uso improprio della piattaforma.",
    )
    add_body(
        doc,
        "È opportuno che il testo finale concordi un tetto massimo di responsabilità economica, ad esempio pari ai corrispettivi versati dal Cliente in un determinato periodo antecedente all'evento dannoso, fatto salvo quanto non limitabile per legge.",
    )

    add_heading(doc, "15. Durata, rinnovo e recesso", 1)
    add_body(
        doc,
        "Il contratto avrà durata iniziale pari a [durata], con eventuale rinnovo secondo quanto indicato nelle condizioni economiche. Le parti potranno disciplinare recesso ordinario, recesso per giusta causa e termini di preavviso nell'appendice economica o nel testo definitivo.",
    )

    add_heading(doc, "16. Fine rapporto, export dati e disattivazione", 1)
    add_bullets(
        doc,
        [
            "Alla cessazione del rapporto il Cliente può richiedere l'export dei propri dati nei limiti tecnicamente disponibili e secondo formati ragionevolmente utilizzati dal servizio.",
            "Il Fornitore potrà prevedere tempi tecnici, costi di attività straordinarie ed esclusioni per personalizzazioni non standard o dati derivanti da terze parti.",
            "Decorso il periodo concordato per il recupero dati, il Fornitore potrà procedere alla disattivazione degli accessi e alla successiva cancellazione o anonimizzazione dei dati, fatte salve esigenze di legge, sicurezza o conservazione tecnica temporanea.",
        ],
    )

    add_heading(doc, "17. Legge applicabile e foro", 1)
    add_body(
        doc,
        "Il contratto è regolato dalla legge italiana. Il foro competente sarà individuato nel testo definitivo, tenendo conto della natura delle parti e dell'eventuale inderogabilità prevista dalla legge.",
    )

    add_heading(doc, "18. Allegati", 1)
    add_bullets(
        doc,
        [
            "Allegato A - Data Processing Agreement / Allegato Privacy",
            "Allegato B - Allegato tecnico infrastrutturale e funzionale",
            "Allegato C - Offerta economica / condizioni commerciali (da predisporre)",
        ],
    )

    add_page_break_section(doc)
    add_footer(doc.sections[-1])
    add_title(doc, "Allegato A - Data Processing Agreement (bozza)")
    add_subtitle(doc, "Schema privacy da adattare con il consulente legale/privacy del progetto.")

    add_heading(doc, "A.1 Ruoli privacy", 1)
    add_body(
        doc,
        "Nell'ambito dell'erogazione del servizio StoreHub, il Cliente opera di regola quale Titolare del trattamento dei dati personali trattati per le proprie finalità organizzative e operative; il Fornitore opera di regola quale Responsabile del trattamento ai sensi dell'art. 28 GDPR, limitatamente alle attività svolte per conto del Cliente.",
    )

    add_heading(doc, "A.2 Oggetto, natura e finalità del trattamento", 1)
    add_bullets(
        doc,
        [
            "hosting applicativo e messa a disposizione della piattaforma;",
            "gestione utenti, autenticazione, autorizzazioni e tenant;",
            "memorizzazione e consultazione di dati operativi dei punti vendita;",
            "gestione log, sicurezza, backup, supporto e manutenzione;",
            "erogazione di moduli opzionali e integrazioni espressamente attivate dal Cliente.",
        ],
    )

    add_heading(doc, "A.3 Categorie di dati e interessati", 1)
    add_bullets(
        doc,
        [
            "dati identificativi e di contatto di utenti interni del Cliente;",
            "dati organizzativi e operativi relativi a store, personale, orari, survey, report, rendiconti e processi amministrativi;",
            "eventuali immagini, allegati, note operative o dati caricati dal Cliente nei moduli attivati;",
            "log tecnici, credenziali applicative, audit trail e metadati di utilizzo.",
        ],
    )

    add_heading(doc, "A.4 Istruzioni e obblighi del Responsabile", 1)
    add_numbered(
        doc,
        [
            "trattare i dati solo su istruzione documentata del Cliente e nei limiti del servizio contrattualizzato;",
            "garantire che le persone autorizzate al trattamento siano vincolate alla riservatezza;",
            "adottare misure tecniche e organizzative adeguate ai sensi dell'art. 32 GDPR;",
            "assistere il Cliente, nei limiti del servizio e delle informazioni disponibili, nella gestione dei diritti degli interessati, delle valutazioni di impatto e degli incidenti di sicurezza;",
            "cancellare o restituire i dati al termine del rapporto secondo quanto pattuito, salvo obblighi di conservazione.",
        ],
    )

    add_heading(doc, "A.5 Sub-responsabili e fornitori terzi", 1)
    add_body(
        doc,
        "Il Cliente autorizza il Fornitore a utilizzare subfornitori e infrastrutture terze strettamente necessari all'erogazione del servizio, inclusi servizi cloud, servizi di autenticazione, servizi di storage, servizi di posta o messaggistica tecnica, a condizione che siano soggetti a obblighi contrattuali adeguati in materia di protezione dei dati.",
    )
    add_body(
        doc,
        "Alla data della presente bozza, rientrano tipicamente nel perimetro tecnico servizi Azure/App Service, SQL Server, Supabase e ulteriori integrazioni eventualmente attivate dal Cliente.",
    )

    add_heading(doc, "A.6 Trasferimenti e localizzazione", 1)
    add_body(
        doc,
        "Per quanto noto allo stato attuale del progetto, il servizio Supabase in uso per StoreHub è collocato in area europea (West EU / Ireland). Eventuali ulteriori localizzazioni, trasferimenti extra-SEE o subfornitori con accessi da Paesi terzi dovranno essere mappati e regolati con adeguate basi di trasferimento, se applicabili.",
    )

    add_heading(doc, "A.7 Misure tecniche e organizzative", 1)
    add_bullets(
        doc,
        [
            "autenticazione utenti e gestione profili/permessi;",
            "segregazione logica per tenant e limitazione degli accessi ai dati;",
            "uso di credenziali e chiavi applicative dedicate per integrazioni esterne;",
            "logging tecnico, monitoraggio e misure di hardening applicativo progressivamente implementate;",
            "backup e procedure di ripristino secondo configurazione infrastrutturale vigente.",
        ],
    )

    add_heading(doc, "A.8 Data breach", 1)
    add_body(
        doc,
        "Il Fornitore informerà il Cliente senza ingiustificato ritardo una volta venuto a conoscenza di una violazione dei dati personali che riguardi i dati trattati per conto del Cliente, mettendo a disposizione le informazioni ragionevolmente disponibili per la gestione dell'evento.",
    )

    add_page_break_section(doc)
    add_footer(doc.sections[-1])
    add_title(doc, "Allegato B - Allegato tecnico infrastrutturale e funzionale")
    add_subtitle(doc, "Riepilogo tecnico sintetico coerente con l'architettura StoreHub oggi nota.")

    add_heading(doc, "B.1 Struttura applicativa", 1)
    add_bullets(
        doc,
        [
            "backend web realizzato in Python / Flask;",
            "pubblicazione dell'applicazione su Azure App Service;",
            "basi dati SQL Server per la componente applicativa e tenant-specifica;",
            "Supabase per autenticazione, profili, alcuni repository e servizi complementari;",
            "possibile compresenza di componenti legacy o ambienti di transizione per tenant specifici.",
        ],
    )

    add_heading(doc, "B.2 Multi-tenant", 1)
    add_body(
        doc,
        "StoreHub supporta una logica multi-tenant. Ogni tenant può avere configurazioni dedicate, moduli attivi/disattivi, limiti operativi specifici, anagrafiche separate e, nei casi previsti, proprio database applicativo SQL dedicato. Il tenant principale può mantenere alcune integrazioni storiche o ibride non necessariamente replicate sui tenant di nuova attivazione.",
    )

    add_heading(doc, "B.3 Moduli standard", 1)
    add_bullets(
        doc,
        [
            "Dashboard",
            "Cruscotto",
            "Magazzino",
            "Rendiconto",
            "Gestione Orari",
            "Link",
            "Estrazioni",
        ],
    )

    add_heading(doc, "B.4 Moduli opzionali / configurabili", 1)
    add_bullets(
        doc,
        [
            "MBO e survey",
            "ordini fornitore e listini multipli",
            "pagine master / configurazione tenant",
            "integrazioni con servizi esterni (es. Microsoft, provider POS, AI, recensioni, survey, provider delivery)",
            "funzioni sperimentali o test abilitate solo su specifici tenant.",
        ],
    )

    add_heading(doc, "B.5 Dati e flussi", 1)
    add_body(
        doc,
        "La piattaforma può trattare dati di anagrafica store, utenti, fornitori, articoli, orari, rendiconti, fotografie, survey, dati di delivery, KPI e ulteriori informazioni operative provenienti da inserimento manuale, import file o integrazioni API. I flussi effettivamente attivi dipendono dai moduli e dalle integrazioni abilitate sul singolo tenant.",
    )

    add_heading(doc, "B.6 Sicurezza e accessi", 1)
    add_bullets(
        doc,
        [
            "accesso autenticato con profili di autorizzazione;",
            "gestione ruoli tenant e, se abilitato, profilo master/piattaforma;",
            "segregazione logica dei tenant e dei relativi store;",
            "possibilità di attivare/disattivare moduli per tenant;",
            "supporto a chiavi API e credenziali per integrazioni esterne.",
        ],
    )

    add_heading(doc, "B.7 Limiti e dipendenze", 1)
    add_body(
        doc,
        "Prestazioni, disponibilità e profondità funzionale di alcuni moduli possono dipendere da qualità dei dati sorgente, configurazioni tenant, servizi cloud di terzi, limiti API, stato della connettività e configurazione dell'infrastruttura SQL / Azure / Supabase vigente al momento dell'erogazione.",
    )

    add_heading(doc, "B.8 Personalizzazioni e governance", 1)
    add_body(
        doc,
        "Le personalizzazioni sviluppate per singoli tenant possono incidere su menu, traduzioni, visibilità moduli, pagine master/admin, strutture dati, flussi di import/export e integrazioni. È opportuno che tali personalizzazioni siano tracciate in ordini di lavoro, appendici o ticket approvati.",
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    build_doc()
    print(str(OUT_PATH))
